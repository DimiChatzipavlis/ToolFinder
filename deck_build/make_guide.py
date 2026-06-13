"""Generate a plain-language presentation guide (speaker map) as a .docx.

Audience note: written so a non-expert classmate could follow it. Jargon is
explained in everyday words. This file is NOT committed to git.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
BLUE = RGBColor(0x47, 0x78, 0xA8)
ORANGE = RGBColor(0xC0, 0x6A, 0x1F)
MUTED = RGBColor(0x55, 0x5F, 0x6B)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def heading(text, size=15, color=NAVY, space_before=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def body(text, italic=False, color=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def slide_block(num, title, seconds, on_screen, say):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"Slide {num} — {title}")
    r.bold = True
    r.font.size = Pt(12.5)
    r.font.color.rgb = BLUE
    t = p.add_run(f"    (~{seconds})")
    t.italic = True
    t.font.size = Pt(10)
    t.font.color.rgb = MUTED

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)
    sr = sp.add_run("On screen: ")
    sr.bold = True
    sp.add_run(on_screen)

    yp = doc.add_paragraph()
    yp.paragraph_format.space_after = Pt(4)
    yr = yp.add_run("Say (in your words): ")
    yr.bold = True
    yp.add_run(say)


# ---------------- title ----------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("ToolFinder — Presentation Guide")
tr.bold = True
tr.font.size = Pt(20)
tr.font.color.rgb = NAVY
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("A plain-language map of the 11 slides — what each shows and what to say")
sr.italic = True
sr.font.size = Pt(11)
sr.font.color.rgb = MUTED

body("Total time: about 10 minutes. Aim for the seconds shown per slide. The big idea in one sentence: "
     "instead of showing an AI every tool it could use, we teach a small model to quickly pick the few tools that fit the request — "
     "and we prove the result honestly.", space_after=8)

# ---------------- glossary ----------------
heading("Words you'll say — in plain terms", size=14)
glossary = [
    ("Tool / MCP tool", "a function an AI assistant can call (e.g. \"create a branch\", \"search issues\"). MCP is just the standard way these tools are described."),
    ("Routing / tool selection", "choosing the right tool for a request before the AI runs it."),
    ("Bi-encoder (our Model A)", "a small model that turns any text into a list of numbers — a 'fingerprint'. We compare the request's fingerprint to each tool's fingerprint and pick the closest. It's fast because every tool's fingerprint is computed once, up front."),
    ("Cross-encoder reranker (Model B)", "a second, slower model that re-reads the request together with each top candidate to double-check the order. Accurate per pair, but it must re-run for every candidate."),
    ("Fine-tuning", "taking a general model and training it a bit more on our specific task. 'Frozen' means we did NOT train it — used as-is."),
    ("Recall@1", "how often the system's single top pick is the correct tool. 1.00 = perfect, higher is better."),
    ("ROC-AUC", "a 0-to-1 score for how well the system tells 'I have a tool for this' apart from 'this isn't my job'. 1.0 = perfect, 0.5 = a coin flip."),
    ("Leakage", "when the test set secretly contains near-copies of the training questions. A model can then score high by memorizing, not by understanding — so the score looks better than it really is. Catching and removing this is our main contribution."),
    ("BM25 / TF-IDF", "classic keyword-matching search (no AI). We use them as honest baselines to beat."),
    ("Flat vs HNSW", "two ways to search the fingerprints. 'Flat' checks every tool exactly; 'HNSW' is an approximate shortcut for huge catalogs."),
]
for term, definition in glossary:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{term}: ")
    r.bold = True
    p.add_run(definition)

# ---------------- slide map ----------------
heading("Slide-by-slide map", size=14, space_before=12)

slide_block(1, "Title", "20s",
            "Project title and your name.",
            "Introduce yourself and the one-line idea: helping AI assistants pick the right tool from a big catalog, reliably.")
slide_block(2, "The problem", "60s",
            "Three problems (with icons) and a dark box on the right.",
            "If you hand an AI every tool at once, the list is too long: it gets confused, picks wrong tools, and small models break. "
            "And picking a wrong tool isn't just annoying — if it's a destructive tool, it's dangerous. So good selection matters.")
slide_block(3, "What others did", "60s",
            "Two columns of prior systems, and an orange takeaway bar.",
            "Others built tools for this. We are NOT claiming a brand-new design — we're claiming we tested it far more honestly than usual. "
            "Say that openly; it makes you credible.")
slide_block(4, "Our benchmark and a hidden flaw", "75s",
            "A histogram, two big numbers (1,695 and 574), and a dark box saying 96%.",
            "We built a test set of 1,695 real requests over 574 real tools. Then we found a trap: with the usual random split, a dumb method that just "
            "memorizes training questions already scores 96%. That means the old way of testing was basically cheating. The orange and blue bars show how "
            "much the test overlaps the training data.")
slide_block(5, "How we test fairly", "60s",
            "Four cards: Regime 1, 1b, 2, 3.",
            "To stop the cheating, we test in four harder, fairer ways: new wordings, brand-new wordings AND scenarios, completely new tools, and even new "
            "servers. A automatic check blocks any overlap. This is the heart of the project.")
slide_block(6, "The two models we trained", "60s",
            "Two cards: Model A (bi-encoder) and Model B (reranker).",
            "We trained two different kinds of AI model. Model A is the fast 'fingerprint' matcher. Model B is the slower double-checker. "
            "We compare them fairly — same data, three random seeds each.")
slide_block(7, "Main result", "75s",
            "A bar chart of all systems, plus two highlight boxes.",
            "Our fine-tuned model wins everywhere — 0.99, 0.91, 0.67 across the three tests, versus about 0.57, 0.76, 0.49 for keyword search. "
            "The surprise: training a TINY model matters far more than using a big one. The untrained version actually loses to plain keyword search.")
slide_block(8, "Proof it's not cheating", "45s",
            "One highlight box (0.958 vs 0.650) and the training curves.",
            "Remember the cheating trap? On the hardest, cheat-proof test our model still scores 0.958, while the memorizing trick drops to 0.650. "
            "So the win is real, not memorized. The curves show training went smoothly.")
slide_block(9, "Safety and speed", "75s",
            "Two highlight boxes (0.99 and 0.0% vs 83%) and a scaling chart.",
            "Two safety wins: it correctly refuses out-of-scope requests (0.99 score), and when a fake malicious tool tries to hijack it, our model is fooled "
            "0% of the time versus 83% for keyword search. On speed: the simple exact search is the right choice — the slow part is reading the request, not the search.")
slide_block(10, "What did NOT work (honesty slide)", "45s",
            "Three cards and a bar chart comparing Model A and Model B.",
            "We report failures too. The fancy double-checker (Model B, orange bars) actually does slightly WORSE than the simple model (blue bars) in every test, "
            "and costs ~140x more. We also note our shipped default isn't the best model, and our data is synthetic. Saying this out loud earns trust.")
slide_block(11, "Conclusion", "45s",
            "Four takeaways and a limitations box.",
            "Wrap up: a small trained model solves this well and safely; training beats size; we honestly report what failed; and everything is reproducible. "
            "End on the limitations so you control the first question.")

# ---------------- likely questions ----------------
heading("If they ask… (plain answers)", size=14, space_before=12)
qa = [
    ("\"Why is one score a perfect 1.00? Isn't that suspicious?\"",
     "Good eye — that's exactly the cheating trap. That perfect score is on the easiest test. On the cheat-proof test it drops to 0.96, which is the honest number."),
    ("\"Why did the second model make things worse?\"",
     "The first model was already almost perfect, so there was little to fix and more to break. The double-checker was trained on very little data, so it added new mistakes. It's an honest negative result."),
    ("\"Is the architecture new?\"",
     "No, and we say so. The contribution is the fair testing method and the measured findings, not a new design."),
    ("\"Why didn't you compare against a big LLM choosing the tool itself?\"",
     "We wrote that experiment but couldn't run it on our machine (it needs a local AI server). It's listed as a limitation, not hidden."),
    ("\"Real users or made-up questions?\"",
     "Made-up but realistic, generated from real tool descriptions. We list 'synthetic data' as our main limitation."),
]
for q, a in qa:
    qp = doc.add_paragraph()
    qp.paragraph_format.space_after = Pt(1)
    qr = qp.add_run(q)
    qr.bold = True
    qr.font.color.rgb = NAVY
    ap = doc.add_paragraph()
    ap.paragraph_format.space_after = Pt(5)
    ap.add_run(a)

heading("One tip", size=14, space_before=12)
body("Your strongest, most original story is the 'cheating trap' (leakage) on slides 4, 8, and 10. If you're short on time, protect those three — "
     "they're what makes the project stand out.", italic=True)

out = Path(__file__).resolve().parents[1] / "reports" / "ToolFinder_presentation_guide.docx"
doc.save(str(out))
print("wrote", out)
